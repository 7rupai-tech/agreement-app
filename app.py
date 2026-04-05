import streamlit as st
from datetime import datetime, timedelta
from docx import Document

# -------------------------------
# PAGE CONFIG
# -------------------------------
st.set_page_config(page_title="Agreement System", layout="wide")

st.title("📄 Client Agreement System")

# -------------------------------
# INPUT SECTION
# -------------------------------
name = st.text_input("Client Name")
amount = st.text_input("Amount (₹)")
date = st.date_input("Investment Date")

lender = st.text_input("Lender Name")
bank = st.text_input("Bank Account")
mode = st.text_input("Payment Mode")
txn = st.text_input("Transaction ID")
email = st.text_input("Email")
phone = st.text_input("Phone Number")
nominee = st.text_input("Nominee Name")
relation = st.text_input("Relation")
fees = st.text_input("Processing Fees")

# -------------------------------
# BUTTON
# -------------------------------
generate = st.button("🚀 Generate Agreement")

# -------------------------------
# FUNCTION: PROCESS DATA
# -------------------------------
def process_data(data):
    investment_date = datetime.strptime(data["INVESTMENT_DATE"], "%d/%m/%Y")

    # 30 days grace
    start_date = investment_date + timedelta(days=30)

    # End date logic
    if investment_date.year == 2025:
        end_date = datetime.strptime("31/12/2026", "%d/%m/%Y")
    else:
        end_date = start_date.replace(year=start_date.year + 1) - timedelta(days=1)

    total_days = (end_date - start_date).days
    months = total_days // 30
    days = total_days % 30

    data["START_DATE"] = start_date.strftime("%d/%m/%Y")
    data["END_DATE"] = end_date.strftime("%d/%m/%Y")
    data["DURATION"] = f"{months} Months {days} Days"

    return data

# -------------------------------
# FUNCTION: GENERATE DOC
# -------------------------------
def generate_doc(data):
    doc = Document()

    doc.add_heading("Loan Agreement", 0)

    doc.add_paragraph(f"Client Name: {data['NAME']}")
    doc.add_paragraph(f"Loan Amount: ₹{data['AMOUNT']}")
    doc.add_paragraph(f"Investment Date: {data['INVESTMENT_DATE']}")
    doc.add_paragraph(f"Start Date: {data['START_DATE']}")
    doc.add_paragraph(f"End Date: {data['END_DATE']}")
    doc.add_paragraph(f"Duration: {data['DURATION']}")

    doc.add_paragraph(f"Lender Name: {data['LENDER_NAME']}")
    doc.add_paragraph(f"Bank: {data['BANK']}")
    doc.add_paragraph(f"Mode: {data['MODE']}")
    doc.add_paragraph(f"Transaction ID: {data['TRANSACTION_ID']}")
    doc.add_paragraph(f"Email: {data['EMAIL']}")
    doc.add_paragraph(f"Phone: {data['PHONE']}")
    doc.add_paragraph(f"Nominee: {data['NOMINEE']} ({data['RELATION']})")
    doc.add_paragraph(f"Processing Fees: ₹{data['PROCESSING_FEES']}")

    file_name = "agreement.docx"
    doc.save(file_name)

    return file_name

# -------------------------------
# MAIN LOGIC
# -------------------------------
if generate:

    if not name or not amount:
        st.error("⚠️ Please fill all required fields")
    else:
        data = {
            "NAME": name,
            "AMOUNT": amount,
            "INVESTMENT_DATE": date.strftime("%d/%m/%Y"),

            "LENDER_NAME": lender,
            "BANK": bank,
            "MODE": mode,
            "TRANSACTION_ID": txn,
            "EMAIL": email,
            "PHONE": phone,
            "NOMINEE": nominee,
            "RELATION": relation,
            "PROCESSING_FEES": fees
        }

        data = process_data(data)

        doc_file = generate_doc(data)

        st.success("✅ Agreement Generated Successfully")

        with open(doc_file, "rb") as f:
            st.download_button("📄 Download Agreement", f, file_name=doc_file)
